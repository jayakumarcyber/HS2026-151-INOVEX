import io
import re
import logging
from datetime import datetime, timezone
from typing import List
import docx

from app.schemas.document import ExtractedDocument, ExtractedPage
from app.services.parsers.base_parser import BaseDocumentParser

logger = logging.getLogger(__name__)


class DOCXParser(BaseDocumentParser):
    """Parser for DOCX documents using python-docx."""

    def clean_text(self, raw_text: str) -> str:
        if not raw_text:
            return ""
        cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", raw_text)
        cleaned = re.sub(r"[ \t]+", " ", cleaned)
        cleaned = re.sub(r"\n\s*\n\s*\n+", "\n\n", cleaned)
        return cleaned.strip()

    def parse(self, file_bytes: bytes, filename: str, document_id: str) -> ExtractedDocument:
        doc = docx.Document(io.BytesIO(file_bytes))
        sections: List[ExtractedPage] = []
        current_heading = "Document Header"
        current_paragraphs: List[str] = []
        section_index = 1

        def flush_section():
            nonlocal section_index, current_heading, current_paragraphs
            text_content = self.clean_text("\n\n".join(current_paragraphs))
            if text_content:
                sections.append(
                    ExtractedPage(
                        document_id=document_id,
                        filename=filename,
                        file_type="docx",
                        page=section_index,
                        section_label=current_heading,
                        text=text_content,
                        char_count=len(text_content),
                    )
                )
                section_index += 1
            current_paragraphs = []

        for p in doc.paragraphs:
            p_text = p.text.strip()
            if not p_text:
                continue

            # Check if paragraph is a heading style
            if p.style and p.style.name and p.style.name.startswith("Heading"):
                flush_section()
                current_heading = p_text
                current_paragraphs.append(f"[{p_text}]")
            else:
                current_paragraphs.append(p_text)

        # Parse tables in docx
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                current_paragraphs.append("\n" + "\n".join(table_rows))

        flush_section()

        if not sections:
            sections.append(
                ExtractedPage(
                    document_id=document_id,
                    filename=filename,
                    file_type="docx",
                    page=1,
                    section_label="General Section",
                    text="[Empty DOCX Document]",
                    char_count=21,
                )
            )

        return ExtractedDocument(
            document_id=document_id,
            filename=filename,
            file_type="docx",
            total_pages=len(sections),
            extracted_at=datetime.now(timezone.utc).isoformat(),
            pages=sections,
        )
