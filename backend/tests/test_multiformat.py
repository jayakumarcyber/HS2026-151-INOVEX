import io
import pytest
import docx
from app.services.parsers import get_parser, get_file_type
from app.services.parsers.pdf_parser import PDFParser
from app.services.parsers.docx_parser import DOCXParser
from app.services.parsers.txt_parser import TXTParser
from app.services.parsers.csv_parser import CSVParser
from app.services.parsers.json_parser import JSONParser
from app.services.parsers.md_parser import MarkdownParser
from app.services.document_processor import document_processor


def test_file_type_validation():
    """Verify supported formats and unsafe format rejections."""
    assert get_file_type("sample.pdf") == "pdf"
    assert get_file_type("sample.docx") == "docx"
    assert get_file_type("sample.txt") == "txt"
    assert get_file_type("sample.csv") == "csv"
    assert get_file_type("sample.json") == "json"
    assert get_file_type("sample.md") == "md"

    # Test unsafe file types
    for unsafe in ["hack.exe", "script.bat", "run.sh", "code.py", "malicious.js"]:
        with pytest.raises(ValueError) as excinfo:
            get_file_type(unsafe)
        assert "Unsupported file type" in str(excinfo.value)


def test_txt_parser():
    """Verify TXT document parsing."""
    parser = TXTParser()
    content = b"Attendance Policy:\nMinimum attendance requirement is 75% per semester."
    extracted = parser.parse(content, "policy.txt", "doc-txt-1")
    assert extracted.file_type == "txt"
    assert len(extracted.pages) == 1
    assert "75%" in extracted.pages[0].text
    assert extracted.pages[0].section_label == "policy.txt"


def test_csv_parser():
    """Verify CSV document parsing and row key-value formatting."""
    parser = CSVParser()
    csv_bytes = (
        "question,answer\n"
        "What is attendance?,75% minimum attendance required\n"
        "How many books can be borrowed?,Students can borrow up to 3 books\n"
    ).encode("utf-8")

    extracted = parser.parse(csv_bytes, "faq.csv", "doc-csv-1")
    assert extracted.file_type == "csv"
    assert len(extracted.pages) == 2
    assert "Question: What is attendance?" in extracted.pages[0].text
    assert "Answer: 75% minimum attendance required" in extracted.pages[0].text
    assert extracted.pages[0].section_label == "Row 1"
    assert extracted.pages[1].section_label == "Row 2"


def test_json_parser():
    """Verify JSON document parsing and flattening, and invalid JSON error handling."""
    parser = JSONParser()
    json_bytes = b'{"college": {"attendance": "75%", "library": {"max_books": 3}}}'
    extracted = parser.parse(json_bytes, "info.json", "doc-json-1")
    assert extracted.file_type == "json"
    assert "college.attendance: 75%" in extracted.pages[0].text

    # Test invalid JSON handling
    invalid_bytes = b'{"college": {"attendance": "75%"' # Missing closing brace
    with pytest.raises(ValueError) as excinfo:
        parser.parse(invalid_bytes, "invalid.json", "doc-json-2")
    assert "Unable to process this JSON file because the file format is invalid." in str(excinfo.value)


def test_md_parser():
    """Verify Markdown parsing preserving heading titles as section labels."""
    parser = MarkdownParser()
    md_bytes = (
        "# Hostel Rules\n\n"
        "Students must return to the hostel by 9:00 PM.\n\n"
        "## Examination Policy\n\n"
        "Mobile phones are strictly prohibited in the exam hall.\n"
    ).encode("utf-8")

    extracted = parser.parse(md_bytes, "rules.md", "doc-md-1")
    assert extracted.file_type == "md"
    assert len(extracted.pages) == 2
    assert extracted.pages[0].section_label == "Hostel Rules"
    assert "9:00 PM" in extracted.pages[0].text
    assert extracted.pages[1].section_label == "Examination Policy"
    assert "Mobile phones" in extracted.pages[1].text


def test_docx_parser():
    """Verify DOCX document parsing with python-docx."""
    doc = docx.Document()
    doc.add_heading("Attendance Requirement", level=1)
    doc.add_paragraph("Students must maintain a minimum attendance of 75%.")

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    parser = DOCXParser()
    extracted = parser.parse(docx_bytes, "handbook.docx", "doc-docx-1")
    assert extracted.file_type == "docx"
    assert len(extracted.pages) >= 1
    assert "75%" in extracted.pages[0].text
